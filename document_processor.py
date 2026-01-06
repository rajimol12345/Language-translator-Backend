import requests
import logging
import html
import warnings
from pathlib import Path
from typing import List, Union, Dict

from docx import Document
from ebooklib import epub, ITEM_DOCUMENT
from fpdf import FPDF
from bs4 import BeautifulSoup
import pypdf

# OCR (optional)
try:
    from pdf2image import convert_from_path
    import pytesseract
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

warnings.filterwarnings("ignore", category=UserWarning)
warnings.filterwarnings("ignore", module="pypdf")

logger = logging.getLogger(__name__)
logging.basicConfig(level=logging.INFO)

class DocumentProcessor:
    FONT_URLS = {
        "latin": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSans/NotoSans-Regular.ttf",
        "devanagari": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoSansDevanagari/NotoSansDevanagari-Regular.ttf",
        "arabic": "https://github.com/googlefonts/noto-fonts/raw/main/hinted/ttf/NotoNaskhArabic/NotoNaskhArabic-Regular.ttf",
        "cjk": "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/SimplifiedChinese/NotoSansCJKsc-Regular.otf",
    }

    LANGUAGE_FONT_MAP = {
        "hindi": "devanagari",
        "arabic": "arabic",
        "chinese": "cjk",
        "japanese": "cjk",
        "russian": "latin",
        "spanish": "latin",
        "german": "latin",
        "french": "latin",
        "italian": "latin",
        "portuguese": "latin",
    }

    def __init__(self, output_dir: Union[str, Path]):
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        self.font_dir = Path(__file__).parent / "fonts"
        self.font_dir.mkdir(parents=True, exist_ok=True)
        self.font_cache: Dict[str, Path] = {}

    # -------------------- FONTS --------------------
    def _get_font_for_language(self, language: str) -> Union[Path, None]:
        key = self.LANGUAGE_FONT_MAP.get(language.lower(), "latin")

        if key in self.font_cache:
            return self.font_cache[key]

        url = self.FONT_URLS[key]
        font_path = self.font_dir / Path(url).name

        if not font_path.exists():
            try:
                logger.info(f"Downloading font: {font_path.name}")
                r = requests.get(url, timeout=30)
                r.raise_for_status()
                font_path.write_bytes(r.content)
            except Exception as e:
                logger.warning(f"Failed to download font {font_path.name}: {e}")
                return None

        self.font_cache[key] = font_path
        return font_path

    @staticmethod
    def _clean_text(value) -> str:
        if value is None:
            return ""
        text = str(value)
        return "".join(ch for ch in text if ch.isprintable() or ch.isspace()).strip()

    # -------------------- EXTRACTION --------------------
    def extract_paragraphs(self, file_path: str) -> List[Dict[str, str]]:
        path = Path(file_path)
        ext = path.suffix.lower()
        content: List[Dict[str, str]] = []

        if ext == ".docx":
            doc = Document(path)
            for p in doc.paragraphs:
                txt = self._clean_text(p.text)
                if txt:
                    style = getattr(p.style, "name", "Normal")
                    content.append({"text": txt, "style": style})
            return content

        if ext == ".txt":
            return [
                {"text": self._clean_text(l), "style": "Normal"}
                for l in path.read_text(encoding="utf-8", errors="ignore").splitlines()
                if l.strip()
            ]

        if ext == ".pdf":
            reader = pypdf.PdfReader(path)
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    for p in text.split("\n"):
                        if len(p.strip()) > 3:
                            content.append({"text": self._clean_text(p), "style": "Normal"})

            if not content and OCR_AVAILABLE:
                for img in convert_from_path(path, dpi=200):
                    for p in pytesseract.image_to_string(img).split("\n"):
                        if len(p.strip()) > 3:
                            content.append({"text": self._clean_text(p), "style": "Normal"})
            return content

        if ext == ".epub":
            book = epub.read_epub(str(path))
            for item in book.get_items_of_type(ITEM_DOCUMENT):
                soup = BeautifulSoup(item.get_content(), "html.parser")
                for tag in soup.find_all(["p", "h1", "h2", "h3", "li"]):
                    txt = self._clean_text(tag.get_text())
                    if txt:
                        content.append({"text": txt, "style": tag.name})
            return content

        raise ValueError(f"Unsupported format: {ext}")

    # -------------------- EXPORT --------------------
    def save_by_format(self, paragraphs: List[Dict], fmt: str, language: str, job_id: str) -> Path:
        fmt = fmt.lower()
        lang = language.lower()
        target = self.output_dir / f"translated_{lang}_{job_id}.{fmt}"

        if fmt == "pdf":
            self.export_pdf(paragraphs, target, lang)
        elif fmt == "docx":
            self.export_docx(paragraphs, target)
        elif fmt == "epub":
            self.export_epub(paragraphs, target, lang)
        else:
            raise ValueError(f"Invalid output format: {fmt}")

        if not target.exists():
            raise RuntimeError(f"{fmt.upper()} generation failed")
        logger.info(f"Saved {fmt.upper()} → {target}")
        return target

    # -------------------- PDF --------------------
    def export_pdf(self, paragraphs: List[Dict], output_path: Path, language: str):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()

        font_path = self._get_font_for_language(language)
        if font_path:
            try:
                pdf.add_font("Custom", "", str(font_path), uni=True)
                pdf.set_font("Custom", size=11)
            except Exception as e:
                logger.warning(f"Failed to load font for PDF: {e}")
                pdf.set_font("Helvetica", size=11)
        else:
            pdf.set_font("Helvetica", size=11)

        for p in paragraphs:
            size = 16 if "heading" in str(p.get("style", "")).lower() else 11
            pdf.set_font("Custom" if font_path else "Helvetica", size=size)
            pdf.multi_cell(0, 8, self._clean_text(p.get("text", "")))
            pdf.ln(2)

        pdf.output(str(output_path))

    # -------------------- DOCX --------------------
    def export_docx(self, paragraphs: List[Dict], output_path: Path):
        doc = Document()
        for p in paragraphs:
            style = p.get("style", "Normal")
            text = self._clean_text(p.get("text", ""))
            try:
                doc.add_paragraph(text, style=style)
            except Exception:
                doc.add_paragraph(text)
        doc.save(output_path)

    # -------------------- EPUB --------------------
    def export_epub(self, paragraphs: List[Dict], output_path: Path, language: str):
        book = epub.EpubBook()
        book.set_identifier(output_path.stem)
        book.set_title("Translated Document")
        book.set_language(language)

        body = "<html><body>"
        for p in paragraphs:
            body += f"<p>{html.escape(self._clean_text(p.get('text', '')))}</p>"
        body += "</body></html>"

        chapter = epub.EpubHtml(title="Content", file_name="content.xhtml", content=body)
        book.add_item(chapter)
        book.toc = (chapter,)
        book.add_item(epub.EpubNcx())
        book.add_item(epub.EpubNav())
        book.spine = ["nav", chapter]

        epub.write_epub(str(output_path), book)
